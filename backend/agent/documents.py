"""
Document ingestion + retrieval (RAG): users upload files, we chunk and embed them
into a per-session ChromaDB collection, and the agent can search them via a tool.
"""

import os
import re
from io import BytesIO
from typing import Optional

import chromadb
from chromadb.utils.embedding_functions import OpenAIEmbeddingFunction
from pypdf import PdfReader
from docx import Document as DocxDocument

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

_chroma_client: Optional[chromadb.ClientAPI] = None
_collection = None


def _get_collection():
    global _chroma_client, _collection
    if _collection is None:
        _chroma_client = chromadb.PersistentClient(path="./chroma_data")
        ef = OpenAIEmbeddingFunction(
            api_key=os.getenv("OPENAI_API_KEY"),
            model_name="text-embedding-3-small",
        )
        _collection = _chroma_client.get_or_create_collection(
            name="uploaded_documents",
            embedding_function=ef,
        )
    return _collection


def _extract_text(filename: str, raw: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PdfReader(BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if lower.endswith(".docx"):
        doc = DocxDocument(BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return raw.decode("utf-8", errors="ignore")


def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return [c.strip() for c in chunks if c.strip()]


def ingest_document(session_id: str, filename: str, raw: bytes) -> dict:
    """Extract, chunk, and embed a document for later retrieval within this session."""
    text = _extract_text(filename, raw)
    if not text.strip():
        return {"error": "Could not extract any text from this file", "filename": filename, "chunks": 0}

    chunks = _chunk_text(text)
    if not chunks:
        return {"error": "Document produced no usable text chunks", "filename": filename, "chunks": 0}

    col = _get_collection()
    ids = [f"{session_id}_{filename}_{i}" for i in range(len(chunks))]
    metadatas = [{"session_id": session_id, "filename": filename, "chunk_index": i} for i in range(len(chunks))]
    col.upsert(ids=ids, documents=chunks, metadatas=metadatas)

    return {"filename": filename, "chunks": len(chunks), "characters": len(text)}


def search_documents(query: str, session_id: str, n_results: int = 4) -> dict:
    """Search the user's uploaded documents for content relevant to the query."""
    try:
        col = _get_collection()
        results = col.query(
            query_texts=[query],
            n_results=n_results,
            where={"session_id": session_id},
        )
        docs = results.get("documents", [[]])[0]
        metas = results.get("metadatas", [[]])[0]
        if not docs:
            return {"results": [], "note": "No uploaded documents found for this session."}
        return {
            "results": [
                {"filename": meta.get("filename", ""), "excerpt": doc[:800]}
                for doc, meta in zip(docs, metas)
            ]
        }
    except Exception as e:
        return {"error": str(e), "results": []}


def list_documents(session_id: str) -> list[str]:
    try:
        col = _get_collection()
        results = col.get(where={"session_id": session_id})
        filenames = {meta["filename"] for meta in results.get("metadatas", []) if "filename" in meta}
        return sorted(filenames)
    except Exception:
        return []
